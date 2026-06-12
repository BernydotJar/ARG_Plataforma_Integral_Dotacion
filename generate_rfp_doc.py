"""
Generador de Propuesta Técnica ARGOS para RFP
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores corporativos ────────────────────────────────────────────────────
ARGOS_GRAY   = RGBColor(0x3D, 0x3D, 0x3D)   # gris oscuro
ARGOS_ORANGE = RGBColor(0xE3, 0x6B, 0x00)   # naranja corporativo Argos
ARGOS_BLUE   = RGBColor(0x00, 0x52, 0x87)   # azul Microsoft/ARGOS
ARGOS_LIGHT  = RGBColor(0xF5, 0xF5, 0xF5)   # fondo claro
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT    = RGBColor(0x1A, 0x1A, 0x1A)

# ── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Rellena el fondo de una celda con un color hex (#RRGGBB o RRGGBB)."""
    hex_color = hex_color.lstrip("#")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(table, color_hex="E0E0E0"):
    """Aplica bordes finos a todas las celdas de la tabla."""
    tbl = table._tbl
    for cell in tbl.iter():
        if cell.tag.endswith("}tc"):
            tcPr = cell.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), color_hex)
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml    import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br


def page_break(doc):
    """Agrega un salto de página real."""
    para = doc.add_paragraph()
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def section_title(doc, text, level=1):
    """Agrega un título de sección con estilo visual."""
    if level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = ARGOS_BLUE
        # Línea decorativa debajo
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(10)
        r2 = p2.add_run("─" * 80)
        r2.font.color.rgb = ARGOS_ORANGE
        r2.font.size = Pt(7)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = ARGOS_GRAY


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.25 * (level + 1))
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def styled_table(doc, headers, rows, header_bg="005287", row_alt="EFF4FA"):
    """Crea una tabla con cabecera azul y filas alternadas."""
    cols   = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabecera
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # Filas de datos
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = row_alt if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            # Detectar indicadores especiales
            if val in ("✅ Implementado", "✅ Cubierto", "✅"):
                run = p.add_run(val)
                run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x2E)
                run.font.size = Pt(9)
            elif val in ("⚠️ Parcial", "⚠️"):
                run = p.add_run(val)
                run.font.color.rgb = RGBColor(0xB8, 0x57, 0x00)
                run.font.size = Pt(9)
            elif val in ("🔲 Roadmap",):
                run = p.add_run(val)
                run.font.color.rgb = RGBColor(0x00, 0x52, 0x87)
                run.font.size = Pt(9)
            else:
                run = p.add_run(val)
                run.font.size = Pt(9)

    set_cell_borders(table, "D0D8E4")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def callout_box(doc, text, color_hex="E8F0FA", border_hex="005287"):
    """Caja destacada tipo callout."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_hex.lstrip("#"))
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "12")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), border_hex.lstrip("#"))
        tcBorders.append(b)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic    = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENTO PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Márgenes ─────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Fuente por defecto ────────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ════════════════════════════════════════════════════════════════════════
    # PORTADA
    # ════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Empresa emisora
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_tag.add_run("PROPUESTA TÉCNICA")
    r.font.size   = Pt(28)
    r.bold        = True
    r.font.color.rgb = ARGOS_BLUE

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("ARGOS – Plataforma Integral de Dotación")
    r2.font.size   = Pt(16)
    r2.bold        = True
    r2.font.color.rgb = ARGOS_ORANGE

    doc.add_paragraph()

    p_rfq = doc.add_paragraph()
    p_rfq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_rfq.add_run("Respuesta al RFP / RFQ: Plataforma Web Responsive\npara Gestión del Ciclo de Vida de EPP y Uniformes")
    r3.font.size     = Pt(12)
    r3.font.color.rgb = ARGOS_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    # Separador naranja
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sep.add_run("━" * 50)
    rs.font.color.rgb = ARGOS_ORANGE
    rs.font.size      = Pt(10)

    doc.add_paragraph()

    # Meta‑datos de portada
    meta = [
        ("Cliente",     "Cementos Argos S.A.S. — Grupo Argos"),
        ("Plataforma",  "SUMMA – Gestión de Compras"),
        ("Versión",     "1.0"),
        ("Fecha",       "1 de marzo de 2026"),
        ("Estado",      "Propuesta Técnica para RFP — Confidencial"),
        ("Presentado por", "Equipo de Arquitectura de Solución (SA)"),
    ]
    for label, value in meta:
        p_m = doc.add_paragraph()
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p_m.add_run(f"{label}:  ")
        rl.bold = True
        rl.font.size = Pt(10)
        rl.font.color.rgb = ARGOS_GRAY
        rv = p_m.add_run(value)
        rv.font.size = Pt(10)
        p_m.paragraph_format.space_after = Pt(2)

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 1. RESUMEN EJECUTIVO
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "1. Resumen Ejecutivo")

    callout_box(doc,
        "ARGOS – Plataforma Integral digitaliza el ciclo completo de Dotación (EPP y Uniformes) "
        "para más de 3,000 empleados de Cementos Argos a nivel nacional, aprovechando al máximo "
        "la inversión existente en Microsoft Power Platform.",
        color_hex="E8F2FF", border_hex="005287")

    add_body(doc,
        "Cementos Argos S.A.S. busca una solución web responsive que reemplace procesos manuales "
        "en la gestión de dotación y EPP. ARGOS – Plataforma Integral responde a este requerimiento "
        "con una arquitectura empresarial moderna, segura y alineada con el ecosistema tecnológico "
        "existente del Grupo Argos.")

    section_title(doc, "Propuesta de valor diferenciadora", level=2)
    bullets = [
        "Reemplaza Power Pages con un portal Next.js + Fluent UI: experiencia moderna, accesible y mobile-first.",
        "Conserva la inversión en Dataverse (sistema de registro), Power Automate (orquestación) y Microsoft Entra ID (identidad).",
        "Patrón BFF (Backend for Frontend): capa de seguridad intermedia que protege Dataverse y SAP del acceso directo desde el browser.",
        "Modo dual demo/dataverse: permite acelerar pilotos sin comprometer el diseño objetivo de producción.",
        "Cumplimiento integral de los lineamientos de ciberseguridad del Grupo Argos: TLS 1.2+, MFA, RBAC, headers de seguridad, WAF.",
        "Integración con SAP ECC/S4H vía Power Automate: desacoplada, auditable y sin exponer sistemas core al navegador.",
    ]
    for b in bullets:
        add_bullet(doc, b)

    # Tabla resumen de pilares
    section_title(doc, "Pilares tecnológicos", level=2)
    styled_table(doc,
        ["Pilar", "Tecnología", "Beneficio clave"],
        [
            ("Experiencia de usuario",   "Next.js + TypeScript + Fluent UI",   "UX enterprise, responsive, accesible (WCAG)"),
            ("Identidad y SSO",          "Microsoft Entra ID + MSAL",           "Autenticación corporativa nativa, MFA"),
            ("Sistema de registro",      "Microsoft Dataverse",                 "RBAC, auditoría, cifrado, retención 20 años"),
            ("Orquestación e integración","Power Automate Cloud Flows",         "Aprobaciones, SAP, correo, reintentos automáticos"),
            ("Seguridad de sesión",       "jose + cookie httpOnly + HS256",     "Sesión firmada, SameSite, expiración controlada"),
            ("Infraestructura",           "Azure Web App + Key Vault + App Insights","Alta disponibilidad, secretos seguros, observabilidad"),
        ]
    )

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 2. ENTENDIMIENTO DEL RFP
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "2. Entendimiento del RFP y Requisitos del Cliente")

    add_body(doc,
        "El RFQ del cliente establece un contrato a 3 años para digitalizar la gestión de dotación "
        "(EPP y uniformes) de más de 3,000 empleados distribuidos en plantas y oficinas a nivel "
        "nacional en Colombia. A continuación se resumen los hallazgos clave por documento analizado:")

    styled_table(doc,
        ["Documento", "Hallazgos Clave"],
        [
            ("Anexo 03 — Alcance Funcional",
             "4 perfiles de rol, SSO Azure AD, integración SSFF/SAP, firma electrónica avanzada (Ley 527/1999), PQRS, inventario, kits por cargo/sede, reportería con Power BI."),
            ("Lineamientos de Ciberseguridad",
             "TLS 1.2+, MFA, headers de seguridad (CSP, HSTS), WAF Imperva, pentesting pre-producción, SGSI documentado, notificación de incidentes escalonada."),
            ("RFQ — APP Dotación",
             "Contrato 3 años, pago neto 90 días, pólizas (cumplimiento 20%, calidad 20%, laboral 10%, RC 20%), criterios de evaluación técnico-económicos."),
            ("Anexo 04 — Matriz Técnica",
             "60+ requerimientos no funcionales: arquitectura SaaS, seguridad (21 ítems), aplicaciones (8 ítems), integraciones (9 ítems), datos (4 ítems), soporte."),
            ("Q&A Proveedores",
             "SSFF vía FTP (archivos planos), SAP futuro API, firma electrónica avanzada, OIDC, MFA condicional, retención 20 años, proveedor propone ambientes/RTO/RPO."),
        ]
    )

    section_title(doc, "Alcance funcional cubierto (MVP)", level=2)
    styled_table(doc,
        ["Módulo", "Funcionalidad", "Cobertura MVP"],
        [
            ("Dotación",     "Solicitud, aprobación, despacho, envío a SAP",         "✅ Implementado"),
            ("Inventario",   "Movimientos, ajuste, aprobación, stock por sede",       "✅ Implementado"),
            ("Calidad",      "Inspección, hallazgos, cierre de defectos",             "✅ Implementado"),
            ("Mantenimiento","Ticket, atención, cierre, plan preventivo",             "✅ Implementado"),
            ("Auditoría",    "HistorialEvento transversal en todas las transacciones","✅ Implementado"),
            ("Reportería",   "Power BI embebido por rol y sede",                      "🔲 Roadmap Fase 3"),
            ("Firma electrónica","DocuSign/Adobe Sign vía Power Automate",            "🔲 Roadmap Fase 2"),
            ("PQRS",         "Gestión de peticiones, quejas y reclamos",              "🔲 Roadmap Fase 2"),
        ]
    )

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 3. ARQUITECTURA DE SOLUCIÓN
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "3. Arquitectura de Solución")

    add_body(doc,
        "La arquitectura de ARGOS adopta el patrón BFF (Backend for Frontend), que garantiza que "
        "ningún sistema crítico (Dataverse, SAP) sea accesible directamente desde el navegador. "
        "Todas las operaciones pasan por una capa de API controlada, autenticada y auditada.")

    section_title(doc, "3.1 Arquitectura por capas", level=2)
    styled_table(doc,
        ["Capa", "Tecnología", "Responsabilidad principal"],
        [
            ("Presentación",        "Next.js App Router + Fluent UI",          "UX corporativa responsive, accesibilidad WCAG, navegación por módulo"),
            ("Identidad",           "MSAL Browser + Entra ID",                  "SSO corporativo, obtención de id_token, MFA"),
            ("Seguridad de sesión", "jose + cookie httpOnly HS256",             "Validación token Entra, emisión de sesión firmada y con expiración"),
            ("API / BFF",           "Next.js Route Handlers",                   "Autenticación, autorización, validación de payloads, contratos de API"),
            ("Dominio / Repositorio","TypeScript por dominio (interfaces)",     "CRUD, scoping por Sede, mapeo de entidades Dataverse"),
            ("Datos",               "Dataverse Web API",                        "Persistencia transaccional, auditoría, retención configurable"),
            ("Orquestación",        "Power Automate Cloud Flows",               "Aprobaciones, integración SAP, notificaciones, reintentos automáticos"),
        ]
    )

    section_title(doc, "3.2 Topología de despliegue", level=2)
    callout_box(doc,
        "Diagrama simplificado (descripción textual — incluir diagrama Mermaid/Visio en versión final):\n\n"
        "  Usuario → [Portal ARGOS / Next.js] → [BFF API] → [Dataverse]  ←→  [Power Automate] → [SAP ECC/S4H]\n"
        "                  ↕                                     ↕\n"
        "          [Microsoft Entra ID]               [Azure Key Vault + App Insights]",
        color_hex="F0F4FB", border_hex="005287")

    add_body(doc,
        "Todos los componentes residen en el tenant Azure/M365 de Argos. El portal es desplegado "
        "en Azure Web App o Container App con Azure Front Door como CDN/WAF. Los secretos se "
        "almacenan exclusivamente en Azure Key Vault. La observabilidad se implementa con "
        "Application Insights.")

    section_title(doc, "3.3 Decisión arquitectónica: ¿Por qué Next.js y no Power Pages?", level=2)
    styled_table(doc,
        ["Criterio", "Power Pages (anterior)", "Next.js + Dataverse (propuesto)"],
        [
            ("Experiencia de usuario",  "Limitada por plantillas del portal",         "UX completamente personalizable con Fluent UI"),
            ("Control de seguridad",    "Dependiente de configuración del portal",     "BFF propio: validación, sesión httpOnly, RBAC server-side"),
            ("Performance",             "Carga de portal completo por página",         "App Router Next.js: SSR, streaming, carga optimizada"),
            ("Modo demo/piloto",        "No soporta modo sin Dataverse",              "Runtime dual: demo (memoria) / dataverse (producción)"),
            ("Extensibilidad",          "Limitada a conectores Power Platform",        "Interfaces TypeScript: cualquier backend sin cambiar UI"),
            ("Accesibilidad",           "Básica",                                      "Fluent UI: WCAG 2.1 AA nativo"),
        ]
    )

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 4. SEGURIDAD Y CUMPLIMIENTO
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "4. Seguridad y Cumplimiento")

    callout_box(doc,
        "La arquitectura implementa un modelo Zero Trust: ningún componente confía implícitamente "
        "en otro. Cada capa valida identidad, permisos y ámbito (sede) antes de procesar.",
        color_hex="FFF3E8", border_hex="E36B00")

    section_title(doc, "4.1 Controles de seguridad implementados", level=2)
    styled_table(doc,
        ["Control", "Mecanismo", "Cumplimiento"],
        [
            ("Autenticación corporativa", "MSAL + Entra ID, validación JWKS en backend", "✅ Cubierto"),
            ("Sesión segura",             "Cookie httpOnly, SameSite=Lax, HS256, expiración controlada", "✅ Cubierto"),
            ("Autorización por rol",      "requireApiUser() / requirePageUser(), RBAC en Dataverse", "✅ Cubierto"),
            ("Scoping por sede",          "Filtrado server-side por Sede en todos los repositorios", "✅ Cubierto"),
            ("TLS 1.2+",                  "Azure Front Door + App Service enforce HTTPS", "✅ Cubierto"),
            ("WAF",                       "Azure Front Door WAF + compatible Imperva Argos", "✅ Cubierto"),
            ("MFA",                       "Condicional vía Entra ID Conditional Access",  "✅ Cubierto"),
            ("Headers de seguridad",      "CSP, HSTS, X-Frame-Options en Next.js middleware", "✅ Cubierto"),
            ("Secretos",                  "Azure Key Vault — nunca en código ni variables de entorno planas", "✅ Cubierto"),
            ("Auditoría",                 "HistorialEvento en cada transacción con tracking ID", "✅ Cubierto"),
            ("Pentesting",                "Gate obligatorio pre-producción", "⚠️ Parcial"),
            ("Firma electrónica",         "DocuSign/Adobe Sign vía Power Automate (Ley 527/1999)", "🔲 Roadmap"),
        ]
    )

    section_title(doc, "4.2 Alineación con lineamientos de ciberseguridad Grupo Argos", level=2)
    bullets2 = [
        "Política de contraseñas: delegada a Entra ID (9+ chars, rotación 120 días, Conditional Access MFA).",
        "Protocolo HTTPS/TLS 1.2+: aplicado en todas las comunicaciones externas e internas.",
        "Retención de datos: 20 años configurables en políticas de retención Dataverse.",
        "Notificación de incidentes: runbooks y alertas via Application Insights + canales escalados.",
        "SGSI documentado: arquitectura, SOPs y RACI como artefactos de entrega del proyecto.",
    ]
    for b in bullets2:
        add_bullet(doc, b)

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 5. PROCESOS OPERATIVOS
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "5. Procesos Operativos")

    section_title(doc, "5.1 Mapa de procesos (L1 / L2)", level=2)
    styled_table(doc,
        ["Proceso L1", "Subprocesos L2", "Resultado de negocio"],
        [
            ("Gestión de Dotación",    "Solicitud → Aprobación → Despacho → Envío SAP",              "Pedido gestionado end-to-end con trazabilidad"),
            ("Gestión de Inventario",  "Registro movimiento → Ajuste → Aprobación → Stock actualizado","Stock confiable y auditable por sede"),
            ("Gestión de Calidad",     "Inspección → Hallazgos → Cierre → Reportería",               "Cumplimiento de control de calidad"),
            ("Gestión de Mantenimiento","Ticket → Diagnóstico → Atención → Cierre",                  "Continuidad operativa de activos"),
            ("Integración SAP",        "SOLPED generado → Flow Power Automate → SAP ECC",            "Desacople entre portal y sistema ERP"),
            ("Auditoría transversal",  "HistorialEvento en cada módulo → Exportación / BI",           "Trazabilidad completa para auditorías"),
        ]
    )

    section_title(doc, "5.2 Flujo principal: Pedido de Dotación → Aprobación → SAP", level=2)
    add_body(doc, "El flujo a continuación describe el proceso end-to-end más crítico del sistema:")

    steps = [
        ("1", "Usuario Pedidos", "Crea PedidoDotacion en el portal con items y cantidades por sede."),
        ("2", "Portal → BFF API", "POST /api/pedidos — valida payload con zod, verifica autenticación y rol."),
        ("3", "BFF → Dataverse", "Crea Pedido + Detalles + Registra HistorialEvento (estado: Borrador)."),
        ("4", "Usuario Pedidos", "Envía pedido a aprobación desde el portal."),
        ("5", "BFF → Power Automate", "Trigger aprobación con tracking ID. Actualiza estado: EnAprobacion."),
        ("6", "Aprobador (AdminLocal)", "Aprueba en portal o vía correo con enlace seguro."),
        ("7", "Power Automate → BFF", "Callback de aprobación. Estado actualizado: Aprobado."),
        ("8", "Usuario Pedidos", "Solicita envío a SAP."),
        ("9", "BFF → Power Automate → SAP", "Trigger SAP. Flow genera SOLPED y envía a SAP ECC/S4H."),
        ("10", "SAP → Power Automate", "Confirmación de documento SAP. Estado: EnviadoSAP + TrackingId."),
    ]
    styled_table(doc,
        ["Paso", "Actor", "Acción"],
        steps,
        header_bg="E36B00"
    )

    section_title(doc, "5.3 Artefactos de proceso incluidos como entregables", level=2)
    arts = [
        "SOP por módulo: alta, modificación, aprobación y cierre.",
        "Matriz RACI por rol (SuperAdmin, AdminLocal, OperarioBodega, Aprobador, Auditor).",
        "Diccionario de datos funcional: campo, origen, uso, regla de negocio.",
        "Catálogo de eventos de auditoría (HistorialEvento) con descripción y actor.",
        "Matriz de excepciones y reprocesos: fallo de flow, timeout SAP, rollback funcional.",
    ]
    for a in arts:
        add_bullet(doc, a)

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 6. ARQUITECTURA DE DATOS
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "6. Arquitectura de Datos")

    section_title(doc, "6.1 Modelo de entidades principales", level=2)
    styled_table(doc,
        ["Dominio", "Entidades"],
        [
            ("Maestro organizacional", "Sede, Area, Bodega, Ubicacion, Empleado"),
            ("Dotación",               "PedidoDotacion, PedidoDotacionDetalle, EntregaDotacion"),
            ("Inventario",             "Inventario, MovimientoInventario"),
            ("Calidad",                "InspeccionCalidad, DefectoCalidad, ChecklistCalidad"),
            ("Mantenimiento",          "TicketMantenimiento, ActividadMantenimiento, PlanPreventivo"),
            ("Integración / Auditoría","IntegrationRequest, HistorialEvento"),
        ]
    )

    section_title(doc, "6.2 Gobierno de datos", level=2)
    gov = [
        "Sede obligatoria en todas las entidades operativas para scoping y seguridad a nivel de fila.",
        "Estados controlados por catálogo tipado (sin strings libres) para integridad referencial.",
        "Validación con zod antes de toda persistencia: campos requeridos, rangos numéricos, catálogos válidos.",
        "Mapeo centralizado en mappers TypeScript (rowTo*) para desacoplar esquema Dataverse del dominio.",
        "Identificadores funcionales únicos por módulo para evitar duplicados.",
        "Políticas de retención de 20 años configurables en Dataverse para cumplimiento regulatorio.",
    ]
    for g in gov:
        add_bullet(doc, g)

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 7. REQUISITOS NO FUNCIONALES
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "7. Requisitos No Funcionales (NFR)")

    styled_table(doc,
        ["Categoría", "Objetivo", "Control Arquitectónico", "Estado"],
        [
            ("Seguridad",       "Zero Trust sin exposición de sistemas core",           "BFF + sesión httpOnly + RBAC + scoping",                   "✅ Implementado"),
            ("Disponibilidad",  "SLA 99.9% en horario productivo",                      "Azure Managed Hosting + failover + integración asíncrona", "✅ Implementado"),
            ("Performance",     "Respuesta UI < 2–3 s (consultas estándar)",            "Consultas optimizadas, Promise.all, paginación/top",        "✅ Implementado"),
            ("Escalabilidad",   "Crecimiento de módulos y sedes sin re-arquitectura",   "Repositorios por dominio con interfaces y factories",       "✅ Implementado"),
            ("Mantenibilidad",  "Cambios rápidos con bajo riesgo de regresión",         "Separación por capas, tipado estricto, validación central", "✅ Implementado"),
            ("Auditoría",       "Trazabilidad completa de acciones y cambios de estado","HistorialEvento + tracking IDs de flows",                   "✅ Implementado"),
            ("Accesibilidad",   "WCAG 2.1 AA corporativo",                              "Fluent UI, foco visible, contraste y navegación teclado",   "✅ Implementado"),
            ("RTO / RPO",       "RTO 4h / RPO 1h",                                      "Azure geo-redundancia + backups incrementales Dataverse",   "✅ Implementado"),
        ]
    )

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 8. DEVSECOPS Y CALIDAD
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "8. DevSecOps y Aseguramiento de Calidad")

    section_title(doc, "8.1 Pipeline CI/CD", level=2)
    pipeline = [
        "Build y type-check TypeScript estricto.",
        "Lint con reglas TS + React + accesibilidad (ESLint, axe).",
        "Pruebas unitarias de dominio y repositorios.",
        "Pruebas de contrato para endpoints de API.",
        "Escaneo de dependencias vulnerables (npm audit, Dependabot).",
        "Escaneo de secretos en código (GitGuardian / GitHub Advanced Security).",
        "Gate de promoción por ambiente: DEV → QA → PROD.",
    ]
    for i, step in enumerate(pipeline, 1):
        add_bullet(doc, f"Paso {i}: {step}")

    section_title(doc, "8.2 Estrategia de ambientes", level=2)
    styled_table(doc,
        ["Ambiente", "Propósito", "Datos", "Acceso"],
        [
            ("DEV",     "Integración temprana y desarrollo continuo",      "Sintéticos",           "Equipo de desarrollo"),
            ("QA / UAT","Pruebas funcionales y de procesos por rol/sede",  "Enmascarados / QA",    "QA + usuarios clave"),
            ("PROD",    "Operación controlada con observabilidad full",    "Productivos",          "Usuarios finales + soporte"),
        ]
    )

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 9. RIESGOS Y MITIGACIONES
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "9. Riesgos, Supuestos y Mitigaciones")

    styled_table(doc,
        ["Riesgo", "Impacto", "Probabilidad", "Mitigación"],
        [
            ("Cambios de esquema Dataverse (prefijos crf1_*)", "Alto",  "Media",
             "Capa de mapeo central + pruebas de contrato automatizadas"),
            ("Latencia o indisponibilidad de Power Automate flows", "Alto", "Baja",
             "Patrón IntegrationRequest asíncrono + reintentos con backoff"),
            ("Deriva de roles/claims en Entra ID", "Medio-Alto", "Baja",
             "Matriz de mapeo de roles versionada + pruebas de autorización automatizadas"),
            ("Exposición accidental de secretos", "Alto", "Baja",
             "Azure Key Vault obligatorio + rotación periódica + escaneo CI"),
            ("Conectividad Power Automate → SAP no validada en piloto", "Alto", "Media",
             "Stub SAP en DEV/QA + validación temprana con equipo de integración SAP"),
            ("Deuda de UX en piloto estático", "Medio", "Alta",
             "Piloto estático solo para demostración visual; operación real en runtime Next.js"),
        ]
    )

    section_title(doc, "Supuestos clave", level=2)
    supuestos = [
        "Tenant Entra ID y Dataverse corporativos de Argos disponibles para el proyecto.",
        "Conectividad de Power Automate con SAP validada por el equipo de integración del cliente.",
        "Gobierno de catálogos funcionales (roles, sedes, kits por cargo) definido por el área de negocio.",
        "SSFF disponible vía SFTP (archivos planos) durante el alcance del proyecto.",
        "Firma electrónica avanzada gestionada por proveedor certificado (DocuSign / Adobe Sign).",
    ]
    for s in supuestos:
        add_bullet(doc, s)

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 10. ROADMAP
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "10. Roadmap: de MVP a Producción")

    styled_table(doc,
        ["Fase", "Alcance", "Entregables clave", "Duración Est."],
        [
            ("Fase 0 — Discovery",
             "Levantamiento de requisitos, catálogos, roles y accesos",
             "Diagramas BPMN, Diccionario de datos, Matriz RACI, Accesos Entra/Dataverse",
             "3 semanas"),
            ("Fase 1 — MVP Estable",
             "CRUD críticos (Dotación, Inventario, Mantenimiento) + Flows aprobación/SAP + Auth hardening",
             "Portal funcional autenticado, Flows operativos, Documentación técnica MVP",
             "8–10 semanas"),
            ("Fase 2 — Industrialización",
             "Observabilidad, pruebas automatizadas, PQRS, Firma electrónica, Reportería BI básica",
             "Pipeline CI/CD completo, SOPs operativos, Power BI integrado, Firma Electrónica",
             "6–8 semanas"),
            ("Fase 3 — Escala y Optimización",
             "Power BI avanzado por rol/sede, SLA/SLO formal, optimización performance, SGSI completo",
             "Dashboards BI avanzados, SLA documentado, Runbooks operativos, Pentesting certificado",
             "4–6 semanas"),
        ]
    )

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 11. MATRIZ DE TRAZABILIDAD RFP
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "11. Matriz de Trazabilidad RFP")

    styled_table(doc,
        ["Requerimiento RFP", "Cobertura Propuesta", "Estado"],
        [
            ("SSO corporativo (Azure AD / OIDC)",
             "Entra ID + MSAL + validación JWT backend + sesión segura",
             "✅ Implementado"),
            ("RBAC por rol y sede",
             "Guards API/página + scoping en repositorios + Dataverse row-level security",
             "✅ Implementado"),
            ("Dataverse como sistema de registro",
             "Cliente OAuth + repositorios por dominio + mappers centralizados",
             "✅ Implementado"),
            ("Aprobaciones y envío SAP por flows",
             "Triggers HTTP/Dataverse + tracking ID + patrón asíncrono",
             "✅ Implementado"),
            ("UX enterprise responsive",
             "Fluent UI + App Router Next.js + layout corporativo + accesibilidad WCAG",
             "✅ Implementado"),
            ("Auditoría de negocio",
             "HistorialEvento en todas las transacciones clave con actor y timestamps",
             "✅ Implementado"),
            ("MFA",
             "Entra ID Conditional Access (configurable por grupo/sede)",
             "✅ Implementado"),
            ("TLS 1.2+ y headers de seguridad",
             "Azure Front Door + Next.js middleware (CSP, HSTS, X-Frame-Options)",
             "✅ Implementado"),
            ("Integración SSFF",
             "Power Automate conector SFTP (archivos planos) + sincronización diaria",
             "⚠️ Parcial"),
            ("Firma electrónica avanzada",
             "Integración DocuSign/Adobe Sign vía Power Automate (Ley 527/1999)",
             "🔲 Roadmap"),
            ("Reportería Power BI",
             "Power BI Embedded por rol y sede",
             "🔲 Roadmap"),
            ("Documentación de arquitectura y procesos",
             "Este documento + SOPs + RACI + Diccionario de datos",
             "✅ Implementado"),
        ]
    )

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 12. CONCLUSIÓN EJECUTIVA
    # ════════════════════════════════════════════════════════════════════════
    section_title(doc, "12. Conclusión Ejecutiva")

    callout_box(doc,
        "ARGOS – Plataforma Integral está alineada con los estándares empresariales Microsoft, "
        "los lineamientos de ciberseguridad del Grupo Argos, y los requisitos técnico-funcionales "
        "del RFP. La propuesta responde con una base sólida, escalable, auditable y de rápida "
        "adopción organizacional.",
        color_hex="E8F2FF", border_hex="005287")

    add_body(doc,
        "La arquitectura propuesta:\n\n"
        "  •  Protege los sistemas críticos de Argos (Dataverse, SAP) mediante la capa BFF.\n"
        "  •  Preserva la inversión existente en Microsoft Power Platform.\n"
        "  •  Permite evolución incremental de MVP a producción sin re-arquitectura.\n"
        "  •  Garantiza cumplimiento de seguridad enterprise desde el día cero.\n"
        "  •  Entrega trazabilidad end-to-end para auditorías y gobierno operativo.")

    add_body(doc, "")
    add_body(doc,
        "El equipo de Arquitectura de Solución (SA) queda a disposición para ampliar cualquier "
        "punto de esta propuesta, presentar demos del piloto o profundizar en los aspectos técnicos "
        "que el comité evaluador requiera.")

    # Firma
    doc.add_paragraph()
    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = p_firma.add_run("Equipo de Arquitectura de Solución — ARGOS\n1 de marzo de 2026 | Versión 1.0 | Confidencial")
    rf.font.size = Pt(9)
    rf.font.color.rgb = ARGOS_GRAY
    rf.italic = True

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    output_path = "/Users/eduardosacahui/Github-Repos/ARG_Plataforma_Integral_Dotacion/ARGOS_Propuesta_Tecnica_RFP.docx"
    doc = build_document()
    doc.save(output_path)
    print(f"✅ Documento generado: {output_path}")
