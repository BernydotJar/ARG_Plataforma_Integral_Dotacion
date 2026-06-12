"""
Propuesta Técnica – ARGOS Dotación | Arquitectura Web Custom + Azure
Genera el documento Word para el RFP de Cementos Argos / Summa
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paleta de colores ────────────────────────────────────────────────────────
C_BLUE   = RGBColor(0x00, 0x52, 0x87)   # Azul corporativo
C_ORANGE = RGBColor(0xE3, 0x6B, 0x00)   # Naranja Argos
C_GRAY   = RGBColor(0x3D, 0x3D, 0x3D)   # Texto gris oscuro
C_GREEN  = RGBColor(0x1A, 0x7A, 0x2E)   # Verde estado OK
C_AMBER  = RGBColor(0xB8, 0x57, 0x00)   # Naranja estado parcial
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_LTBLUE = RGBColor(0xE8, 0xF2, 0xFF)   # Fondo callout azul


# ─── Utilidades de bajo nivel ─────────────────────────────────────────────────

def _shd(cell, hex6: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6.lstrip("#"))
    tcPr.append(shd)


def _borders(table, color="C8D4E8"):
    for cell in table._tbl.iter():
        if not cell.tag.endswith("}tc"):
            continue
        tcPr = cell.get_or_add_tcPr()
        tcB  = OxmlElement("w:tcBorders")
        for s in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{s}")
            b.set(qn("w:val"),   "single")
            b.set(qn("w:sz"),    "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), color)
            tcB.append(b)
        tcPr.append(tcB)


def _page_break(doc):
    p  = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


# ─── Componentes de diseño ────────────────────────────────────────────────────

def h1(doc, text):
    """Encabezado de sección principal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size  = Pt(13)
    r.font.color.rgb = C_BLUE
    # línea decorativa
    ln = doc.add_paragraph()
    ln.paragraph_format.space_before = Pt(0)
    ln.paragraph_format.space_after  = Pt(8)
    rl = ln.add_run("▬" * 75)
    rl.font.size  = Pt(6)
    rl.font.color.rgb = C_ORANGE


def h2(doc, text):
    """Encabezado de subsección."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(11)
    r.font.color.rgb = C_GRAY


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def bullet(doc, text, lvl=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.3 * (lvl + 1))
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10)


def callout(doc, text, bg="E8F2FF", border="005287"):
    t    = doc.add_table(1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    _shd(cell, bg)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for s in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{s}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "14")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), border)
        tcB.append(b)
    tcPr.append(tcB)
    p  = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r  = p.add_run(text)
    r.font.size = Pt(10)
    r.italic    = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table(doc, headers, rows,
          hdr_bg="005287", alt_bg="EFF4FA", text_size=9):
    """Tabla con cabecera azul y filas alternadas."""
    t   = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style     = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabecera
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        _shd(c, hdr_bg)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(h)
        r.bold = True
        r.font.size  = Pt(text_size)
        r.font.color.rgb = C_WHITE

    # Filas
    STATUS_COLORS = {
        "✅": C_GREEN, "⚠️": C_AMBER,
        "✅ SI": C_GREEN, "⚠️ Parcial": C_AMBER,
    }
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            _shd(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(val)
            r.font.size = Pt(text_size)
            # Colorear indicadores de estado
            for key, color in STATUS_COLORS.items():
                if val.startswith(key):
                    r.font.color.rgb = color
                    break

    _borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def divider_orange(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run("─" * 90)
    r.font.size  = Pt(6)
    r.font.color.rgb = C_ORANGE


# ─── PORTADA ──────────────────────────────────────────────────────────────────

def portada(doc):
    for _ in range(4):
        doc.add_paragraph()

    def cp(text, size, color, bold=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold  = bold
        r.italic = italic
        r.font.size  = Pt(size)
        r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(4)
        return p

    cp("PROPUESTA TÉCNICA",          30, C_BLUE,   bold=True)
    cp("Plataforma Web de Gestión de Dotación", 17, C_ORANGE, bold=True)
    cp("Cementos Argos S.A.S. — Grupo Argos",   13, C_GRAY)

    doc.add_paragraph()
    divider_orange(doc)
    doc.add_paragraph()

    cp("Arquitectura:  Portal Web Custom + Azure Services", 11, C_GRAY, bold=True)
    cp("( Next.js Full-Stack  ·  TypeScript  ·  Azure SQL  ·  Microsoft Entra ID )", 10, C_GRAY)

    doc.add_paragraph()

    meta = [
        ("Cliente",          "Cementos Argos S.A.S. — vía Summa Servicios Corporativos"),
        ("Versión",          "3.0 — Arquitectura Web Custom"),
        ("Fecha",            "Marzo 2026"),
        ("Estado",           "Propuesta Técnica para RFP — Confidencial"),
        ("Elaborado por",    "Equipo de Arquitectura de Solución — DATTICS S.A.S."),
    ]
    for lbl, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        rl = p.add_run(f"{lbl}:  ")
        rl.bold = True
        rl.font.size = Pt(10)
        rl.font.color.rgb = C_GRAY
        rv = p.add_run(val)
        rv.font.size = Pt(10)

    _page_break(doc)


# ═════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    portada(doc)

    # ── 1. RESUMEN EJECUTIVO ──────────────────────────────────────────────────
    h1(doc, "1. Resumen Ejecutivo")

    callout(doc,
        "Este proceso impacta a todos los colaboradores de la organización en todas las "
        "plantas/sedes de la regional. Se requiere una nueva solución que habilite funcionalidades "
        "a cada colaborador: actualizar su dotación y tallas, devoluciones, solicitud de garantías; "
        "y que permita al equipo de Servicios y Productividad realizar carga masiva de dotación, "
        "registro de entregas, administración de bodegas e inventario, garantías y generación de "
        "reportes — incluyendo el reporte para SAP (SOLPED).",
        bg="FFF3E0", border="E36B00")

    body(doc,
        "DATTICS presenta una plataforma web 100% personalizada — un Portal Web Custom sobre "
        "servicios Microsoft Azure — que digitaliza el ciclo completo de dotación (EPP y uniformes) "
        "para más de 3,000 colaboradores distribuidos en plantas y oficinas a nivel nacional en Colombia.")

    h2(doc, "¿Por qué Web Custom y no Power Pages?")
    table(doc,
        ["Criterio", "Power Pages + Power Platform", "Web Custom + Azure (propuesta)"],
        [
            ("Costo anual licenciamiento",
             "$48,000 – $96,000 USD/año\n(usuarios autenticados + flujos premium)",
             "$12,720 USD/año\n(infraestructura Azure solamente)"),
            ("Ahorro en 3 años",
             "Hasta $288,000 USD",
             "$38,160 USD  →  Ahorro de hasta $250,000 USD"),
            ("Control lógica de negocio compleja",
             "Limitado por el motor de Power Pages",
             "Control total: kits paramétricos, ventanas de bloqueo, entregas parciales"),
            ("Propiedad del código",
             "Sin código fuente; dependencia perpetua de licencias",
             "Código fuente propiedad de Argos, en su tenant Azure"),
            ("Experiencia de usuario",
             "Plantillas de portal; personalización limitada",
             "UX/UI completamente adaptada al entorno industrial (planta/móvil)"),
            ("Escalabilidad de módulos",
             "Requiere licencias adicionales por módulo",
             "Sin costo adicional de licenciamiento al crecer en módulos"),
            ("Certificaciones de seguridad",
             "ISO 27001, SOC 2 (Microsoft)",
             "Idénticas: ISO 27001, SOC 2, GDPR — heredadas de Azure"),
        ],
        hdr_bg="E36B00"
    )

    _page_break(doc)

    # ── 2. COMPROMISO DE ENTREGA EN 3 MESES ──────────────────────────────────
    h1(doc, "2. Compromiso de Entrega: MVP en 4 Meses · Go-Live Total en 6 Meses")

    callout(doc,
        "DATTICS entrega el MVP completo en producción en 16 semanas (4 meses). "
        "A los 90 días (Sem 13), la plataforma estará funcional y con UAT aprobado: "
        "autenticación, tallas, motor de kits, entrega con firma, SOLPED SAP y reportes. "
        "Solo quedan pentest y Go-Live (Sem 13–16). La Fase 2 se entrega 6 semanas después. "
        "Total proyecto: ~22 semanas.",
        bg="FFF3E0", border="E36B00")

    h2(doc, "2.1 ¿Qué estará disponible y cuándo?")
    table(doc,
        ["Capacidad", "Disponibilidad", "Detalle"],
        [
            ("Login colaboradores y admins",        "✅ Mes 1 (Sem 3–5)",
             "SSO OIDC (23 admins) + Entra B2C cédula/contraseña/MFA (3,000 colaboradores)"),
            ("Sincronización SuccessFactors",       "✅ Mes 1 (Sem 3–4)",
             "Job SFTP automático diario: altas, bajas, cambios de sede/cargo"),
            ("Motor de kits + administración",      "✅ Mes 2 (Sem 5–7)",
             "Asignación automática por cargo/género/sede/ciclo. Super Admin configura kits y maestros."),
            ("Tallas self-service con ventanas",    "✅ Mes 2 (Sem 6–8)",
             "Colaborador actualiza sus tallas con ventanas de bloqueo por sede y nacional"),
            ("Registro de entrega + firma + PDF",   "✅ Mes 2–3 (Sem 8–10)",
             "Admin Local registra entrega. Colaborador firma digitalmente. PDF enviado por correo."),
            ("Generación SOLPED + cargas masivas",  "✅ Mes 3 (Sem 10–11)",
             "Tablero de cierre por sede. Descarga archivo SAP. Cargas masivas Excel/CSV validadas."),
            ("Reportes operativos + Power BI",      "✅ Mes 3 (Sem 11–12)",
             "Colaboradores por sede, dotación, plano de facturación, entregas y Power BI básico"),
            ("UAT aprobado — HITO 90 DÍAS",         "✅ Mes 3 (Sem 12–13)",
             "Pruebas QA, migración de datos y UAT firmado con equipo Argos."),
            ("Pentest + Go-Live MVP",               "✅ Mes 4 (Sem 13–16)",
             "Ethical Hacking, remediación e informe formal. Go-Live con soporte presencial."),
            ("Inventario y bodegas",                "⚠️ Fase 2 — Mes 5 (Sem 17–19)",
             "Módulo completo de bodegas con movimientos entre sedes"),
            ("PQRS + alarmas + proveedores",        "⚠️ Fase 2 — Mes 5–6 (Sem 19–21)",
             "Canal PQRS bidireccional + alarmas configurables + notificaciones a proveedor"),
            ("Go-Live Fase 2",                      "⚠️ Fase 2 — Mes 6 (Sem 21–22)",
             "Reportes avanzados, QA, UAT y Go-Live final de Fase 2."),
        ]
    )

    h2(doc, "2.2 Hitos de entrega — Fase 1 (16 semanas / 4 meses)")
    table(doc,
        ["Semana", "Entregable al cliente", "Hito verificable"],
        [
            ("Sem 1–2",
             "Levantamiento funcional + infraestructura Azure + CI/CD operativo (paralelo)",
             "HU y wireframes revisados con Argos. Ambientes DEV/QA/PROD accesibles."),
            ("Sem 3–5",
             "Autenticación SSO (admins) + B2C (colaboradores) + integración SFTP activa (paralelo)",
             "Demo: login con cuenta Argos y con cédula+contraseña. Job SFTP ejecuta correctamente."),
            ("Sem 5–8",
             "Módulo administración (maestros, kits, sedes) + módulo tallas con ventanas",
             "Demo: Super Admin configura kits. Admin Local gestiona sede. Colaborador actualiza tallas."),
            ("Sem 8–10",
             "Registro de entrega con firma electrónica + generación PDF + notificación por correo",
             "Demo: Admin Local registra entrega. Colaborador firma digitalmente. PDF recibido."),
            ("Sem 10–12",
             "Generación SOLPED para SAP + cargas masivas + reportes operativos + Power BI",
             "Demo: archivo SAP generado. Carga masiva validada. 9 reportes operativos disponibles."),
            ("Sem 12–13  ⭐ HITO 90 DÍAS",
             "Migración de datos históricos + pruebas QA + UAT con equipo Argos",
             "UAT firmado. Datos migrados. Plataforma completa y validada en Sem 13."),
            ("Sem 13–14",
             "Ethical Hacking + remediación de hallazgos",
             "Informe formal de pentest entregado. Hallazgos críticos/altos resueltos."),
            ("Sem 15–16",
             "Go-Live + soporte presencial + estabilización en producción",
             "Sistema en PRODUCCIÓN. Soporte on-site durante estabilización."),
        ],
        hdr_bg="E36B00"
    )

    h2(doc, "2.3 Condiciones para cumplir el cronograma de 4 meses")
    body(doc,
        "El cumplimiento del cronograma de 16 semanas para Fase 1 depende de las siguientes "
        "condiciones por parte de Cementos Argos / Summa:")
    condiciones = [
        "Acceso al tenant Azure AD y Entra ID de Argos en la semana 1 (para configurar OIDC y B2C).",
        "Formato y credenciales SFTP de SuccessFactors disponibles en la semana 1.",
        "Definición completa de los kits de dotación (cargo, género, sede, ciclo) en la semana 1.",
        "Aprobación de wireframes y diseño de UI en máximo 5 días hábiles tras entrega (Sem 2).",
        "Disponibilidad de usuarios clave de Argos para revisiones semanales (30–60 min).",
        "Contrato con proveedor de firma electrónica (DocuSign o Certicamara) gestionado antes de Sem 6.",
        "UAT iniciado en semana 12 con usuarios representativos de cada perfil y sede.",
        "Equipo de datos de Argos disponible en semana 12 para soporte en migración de datos históricos.",
    ]
    for c in condiciones:
        bullet(doc, c)

    _page_break(doc)

    # ── 3. ALCANCE FUNCIONAL ──────────────────────────────────────────────────
    h1(doc, "3. Alcance Funcional — Capacidades del Sistema")

    callout(doc,
        "La solución cubre el ciclo completo de dotación de forma self-service para el colaborador "
        "y con herramientas de gestión avanzada para los administradores.",
        bg="E8F2FF", border="005287")

    h2(doc, "2.1 Perfiles de usuario y sus capacidades")
    table(doc,
        ["Perfil", "Cant.", "Autenticación", "Capacidades principales"],
        [
            ("Super Administrador",  "4",
             "SSO OIDC (Azure AD)",
             "Configuración nacional: kits, roles, sedes, ventanas de bloqueo, maestros, reportes nacionales, cargas masivas globales"),
            ("Administrador Local",  "18",
             "SSO OIDC (Azure AD)",
             "Gestión por sede: entregas con firma, inventario/bodegas, cargas masivas de tallas, movimientos entre sedes, reportes de sede"),
            ("Usuario Pedidos",      "1",
             "SSO OIDC (Azure AD)",
             "Apertura/cierre de períodos, validación de sedes, generación del archivo SOLPED para SAP"),
            ("Colaborador (self-service)", "3,000+",
             "Entra External ID (B2C)\nCédula + contraseña + MFA",
             "Consultar dotación asignada, actualizar tallas, ver historial de entregas, firmar recepción electrónicamente, registrar PQRS y garantías"),
        ]
    )

    h2(doc, "2.2 Módulos — Fase 1 (MVP)")
    table(doc,
        ["Módulo", "Descripción", "Actor principal"],
        [
            ("Integración SuccessFactors",
             "Job automático diario (Azure Functions) vía SFTP: altas, bajas, cambios de sede/cargo, reingresos. Reporte de última sincronización para Admin Local.",
             "Sistema / Super Admin"),
            ("Autenticación dual SSO + Cédula",
             "SSO OIDC para 23 usuarios corporativos. Auth local con Entra External ID para 3,000 colaboradores: cédula + contraseña + MFA. CAPTCHA, bloqueo tras 5 intentos, expiración 120 días.",
             "Todos"),
            ("Carga masiva inicial (migración)",
             "Portal para carga de archivos Excel/CSV con validación: colaboradores, dotación/tallas, sedes, kits, prendas, precios. Auditoría de quién cargó qué y cuándo.",
             "Super Admin"),
            ("Actualización de tallas (self-service)",
             "Colaborador actualiza sus tallas dentro de ventanas parametrizables. Bloqueo por sede (Admin Local) y bloqueo nacional (Super Admin). Admin Local puede hacer cargas masivas de tallas.",
             "Colaborador / Admin Local"),
            ("Motor de kits de dotación",
             "Cálculo automático de kit por género, cargo, sede y ciclo. Se recalcula al cambio de posición/sede (datos de SSFF). Prendas obligatorias y opcionales por sede.",
             "Sistema / Super Admin"),
            ("Registro de entrega + firma electrónica",
             "Admin Local registra prendas entregadas. Colaborador firma electrónicamente (DocuSign/Certicamara, Ley 527/1999). Next.js genera el PDF server-side, lo almacena en Azure Blob Storage y lo envía al correo del colaborador vía Microsoft Graph.",
             "Admin Local / Colaborador"),
            ("Entregas parciales y pendientes",
             "Registro de estado Pendiente con nota. Cierre con evidencia adjunta. Reporte de pendientes por Admin Local.",
             "Admin Local"),
            ("Generación SOLPED para SAP",
             "Tablero de estado por Admin Local. Ajustes por sede. Generación de archivo plano con formato exacto SAP. Futuro: integración directa API/BAPI sin cambiar el portal.",
             "Usuario Pedidos"),
            ("Reportes Fase 1",
             "Colaboradores por sede con dotación, dotación por bodega consolidado nacional, plano de facturación, financiero, entregas por sede, comportamiento mensual. Export Excel/PDF.",
             "Admin Local / Super Admin"),
            ("Administración de maestros",
             "Sedes, prendas con historial de precios, materiales, centros de costo, direcciones de entrega. Gestión completa por Super Admin.",
             "Super Admin"),
        ]
    )

    _page_break(doc)

    h2(doc, "2.3 Módulos — Fase 2 (Funcionalidades Avanzadas)")
    table(doc,
        ["Módulo", "Descripción"],
        [
            ("Bodegas e inventario en línea",
             "Bodegas por sede con inventarios asociados. Consolidado nacional de stock. Movimientos entre sedes con flujo solicitud/aceptación/rechazo y notificaciones."),
            ("Canal PQRS y garantías",
             "Registro de PQRS por colaborador con adjuntos (foto/documentos). Respuestas bidireccionales. Si es calidad/garantía, se genera correo automático al proveedor. Escala de satisfacción al cierre."),
            ("Alarmas y notificaciones",
             "Campana en portal + correo: mínimo stock, falta rotación, solicitudes de movimiento, fechas límite, pedidos sobre cantidad, cambios permanentes de talla."),
            ("Entregas anticipadas",
             "Admin Local registra entregas anticipadas. Se descuenta automáticamente de la próxima entrega periódica."),
            ("Información de proveedores",
             "Registro de proveedores por referencia: NIT, correo, contacto, sedes. Precios asociados. Consulta según rol."),
            ("Log de auditoría completo",
             "Cambio realizado, fecha, usuario, valores old/new, IP. Consultable por Super Admin. Exportable."),
            ("Power BI Pro — Dashboards ANS",
             "Dashboards analíticos avanzados por rol y sede: rotación, inventario, entregas, PQRS. Para 23 administradores."),
        ]
    )

    _page_break(doc)

    # ── 3. ARQUITECTURA TÉCNICA ───────────────────────────────────────────────
    h1(doc, "4. Arquitectura Técnica")

    callout(doc,
        "Arquitectura 100% Azure — sin dependencia de licencias Power Platform. "
        "Todo el stack vive en el tenant de Argos: un único deployment Next.js full-stack "
        "(frontend React + API Routes en el mismo proceso Node.js), que contiene internamente "
        "el motor de kits, el generador SAP y la generación de PDF. "
        "Azure SQL como base de datos, Azure Blob Storage para documentos, y Azure Front Door "
        "como punto de entrada con WAF integrado.",
        bg="E8F2FF", border="005287")

    h2(doc, "3.1 Stack tecnológico por capa")
    table(doc,
        ["Capa", "Componente", "Tecnología", "Hospedaje"],
        [
            ("Presentación + API",
             "Next.js full-stack\n(UI React + API Routes + Server Actions)",
             "Next.js 14+ App Router / TypeScript / Fluent UI React\nUn único deployment: UI, BFF y lógica de negocio.",
             "Azure App Service\n(Node.js 20 LTS — S2+)"),
            ("Integración",
             "Job SFTP — sincronización SuccessFactors",
             "Azure Functions (timer trigger, TypeScript/Node.js)\nRetry automático con backoff exponencial",
             "Azure Functions\n(Consumption Plan)"),
            ("Integración",
             "Firma electrónica avanzada",
             "DocuSign o Certicamara — REST API\nLey 527/1999, no repudio, timestamp",
             "SaaS externo\n(llamado desde Next.js API Route)"),
            ("Integración",
             "Correo — notificaciones y constancias",
             "Microsoft Graph API (O365 del tenant Argos)\nSin dependencia de SendGrid ni Power Automate",
             "Microsoft 365\n(tenant Argos)"),
            ("Datos",
             "Base de datos transaccional",
             "Azure SQL Database + Prisma ORM\nRow-Level Security, TDE AES-256, backups cada 5 min",
             "Azure SQL (S2, 50 DTU\nescalable a S3/S4)"),
            ("Datos",
             "Documentos y evidencias\n(PDFs firmados, fotos de entrega)",
             "Azure Blob Storage\nLifecycle policies: retención configurable hasta 20+ años",
             "Azure Storage\n(LRS o GRS según SLA requerido)"),
            ("Seguridad",
             "Identidad corporativa + colaboradores",
             "Microsoft Entra ID (SSO/OIDC, 23 admins)\n+ Entra External ID B2C (3,000 colaboradores)",
             "Microsoft Cloud\n(tenant Argos)"),
            ("Seguridad",
             "Punto de entrada — CDN + WAF + DDoS",
             "Azure Front Door (WAF con reglas OWASP 3.x)\nDDoS Protection Standard",
             "Azure\n(100% gestionado)"),
            ("Seguridad",
             "Secretos y credenciales",
             "Azure Key Vault\nCero secrets en código ni en variables de entorno",
             "Azure Key Vault"),
            ("Observabilidad",
             "Logs, métricas y alertas",
             "Azure Application Insights + Azure Monitor\nDashboard operativo en portal admin",
             "Azure"),
            ("Analítica",
             "Reportes operativos (export Excel / PDF)",
             "Next.js Server Components + API Routes\nGeneración server-side, sin dependencias externas",
             "Azure App Service"),
        ]
    )

    body(doc,
        "Nota: Motor de kits, generador de SOLPED y generación de PDF son servicios TypeScript "
        "dentro del mismo proceso Next.js. No son deployments adicionales. "
        "Esto simplifica la operación, reduce costos y elimina latencia de red entre componentes.")

    h2(doc, "3.2 Integraciones externas")
    table(doc,
        ["Sistema externo", "Método", "Frecuencia", "Dirección", "Responsable técnico"],
        [
            ("SuccessFactors (SSFF)",
             "SFTP — archivos planos (infodotSSFF/infodot)",
             "Diario automático + retroactivo manual",
             "SSFF → Aplicativo",
             "Azure Functions (job timer, TypeScript/Node.js)"),
            ("Microsoft Entra ID",
             "OIDC / OAuth 2.0 — tokens JWT",
             "Tiempo real (por sesión)",
             "Bidireccional",
             "Next.js middleware — MSAL / NextAuth.js"),
            ("Entra External ID (B2C)",
             "OIDC / OAuth 2.0 — cédula + contraseña + MFA",
             "Tiempo real (por sesión)",
             "Bidireccional",
             "Next.js middleware — MSAL / NextAuth.js"),
            ("SAP ERP",
             "Descarga de archivo plano con layout SOLPED\n(Futuro: API BAPI/OData cuando Argos lo habilite)",
             "Semestral (2× año) — descarga manual por Usuario Pedidos",
             "Aplicativo → SAP",
             "Next.js API Route — generador TypeScript interno"),
            ("Office 365 (correo)",
             "Microsoft Graph API — send mail",
             "Por evento (entrega registrada)",
             "Aplicativo → O365",
             "Next.js API Route — @microsoft/microsoft-graph-client"),
            ("DocuSign / Certicamara",
             "REST API — firma electrónica avanzada",
             "Por evento (entrega con firma)",
             "Bidireccional (solicitud + callback)",
             "Next.js API Route — cliente HTTP (fetch/axios)"),
        ]
    )

    h2(doc, "3.3 Diagrama de arquitectura")
    callout(doc,
        "Colaborador / Admin  (navegador — PC, tablet, celular)\n"
        "         │  HTTPS 443 únicamente\n"
        "         ▼\n"
        "  ┌─────────────────────────────────────────────────────────┐\n"
        "  │   Azure Front Door  (CDN + WAF OWASP 3.x + DDoS)       │\n"
        "  └───────────────────────┬─────────────────────────────────┘\n"
        "                          │\n"
        "  ┌───────────────────────┼───────────────────────────────┐\n"
        "  │         AZURE (tenant Argos)                          │\n"
        "  │                       │                               │\n"
        "  │     [Azure Functions] │                               │\n"
        "  │     Job SFTP (timer)  │                               │\n"
        "  │          │ SFTP       │                               │\n"
        "  │          ▼            ▼                               │\n"
        "  │    [SuccessFactors]  [Next.js Full-Stack App]         │\n"
        "  │                      · React UI (App Router)         │\n"
        "  │                      · API Routes / Server Actions   │\n"
        "  │                      · RBAC + middleware auth        │\n"
        "  │                      · Motor de kits (TypeScript)    │\n"
        "  │                      · Generador SOLPED   [Entra ID] │\n"
        "  │                      · Generación PDF  ◄──► + B2C   │\n"
        "  │                      · Notificaciones               │\n"
        "  │                           │          │               │\n"
        "  │                    [Azure SQL]  [Blob Storage]       │\n"
        "  │                    (Prisma ORM)  (docs/PDFs)         │\n"
        "  │                    [Key Vault]  [App Insights]       │\n"
        "  └──────────────────────────────────────────────────────┘\n"
        "                          │\n"
        "          Llamadas salientes desde Next.js API Routes:\n"
        "     [DocuSign/Certicamara]   [Microsoft Graph/O365]   [SAP — archivo]",
        bg="F0F4FB", border="005287"
    )

    _page_break(doc)

    # ── 4. CAPACIDADES SELF-SERVICE DEL COLABORADOR ──────────────────────────
    h1(doc, "5. Experiencia Self-Service del Colaborador")

    callout(doc,
        "Los 3,000+ colaboradores acceden desde cualquier dispositivo (PC, tablet, smartphone) "
        "sin instalar software. Solo necesitan un navegador web moderno.",
        bg="E8F2FF", border="005287")

    h2(doc, "Jornada del colaborador en la plataforma")
    table(doc,
        ["Paso", "Acción del colaborador", "Lo que hace el sistema"],
        [
            ("1", "Ingresa al portal con su cédula y contraseña",
             "Entra External ID valida credenciales. Si tiene MFA activo, solicita segundo factor. CAPTCHA anti-bot en login."),
            ("2", "Consulta su dotación asignada",
             "Motor de kits calcula y muestra las prendas que le corresponden según su cargo, sede y ciclo actual."),
            ("3", "Actualiza sus tallas (si la ventana está abierta)",
             "Verifica si la ventana de actualización está abierta para su sede. Si sí, guarda las tallas. Si no, muestra el motivo del bloqueo."),
            ("4", "Recibe notificación de entrega disponible",
             "Campana en portal + correo automático vía Microsoft Graph cuando el Admin Local registra una entrega."),
            ("5", "Firma electrónicamente la recepción",
             "Mediante DocuSign/Certicamara: el colaborador firma con validez legal (Ley 527/1999, no repudio, timestamp)."),
            ("6", "Recibe constancia en su correo",
             "PDF generado server-side por Next.js (sin dependencias externas) y enviado al correo del colaborador vía Microsoft Graph. El PDF queda almacenado en Azure Blob Storage con lifecycle policy de retención."),
            ("7", "Registra una PQRS o garantía",
             "Ingresa el reclamo con texto y adjuntos. Si es garantía/calidad, el sistema notifica automáticamente al proveedor (sin datos personales)."),
            ("8", "Consulta su historial",
             "Visualiza todas sus entregas anteriores, certificados descargables, y el estado de sus PQRS."),
        ],
        hdr_bg="E36B00"
    )

    _page_break(doc)

    # ── 5. SEGURIDAD Y CUMPLIMIENTO ───────────────────────────────────────────
    h1(doc, "6. Seguridad y Cumplimiento")

    callout(doc,
        "Cumplimiento total con los Lineamientos de Ciberseguridad del Grupo Argos. "
        "85 de 85 requerimientos técnicos del Anexo 04 cubiertos (100%).",
        bg="FFF3E0", border="E36B00"
    )

    h2(doc, "6.1 Controles de seguridad implementados")
    table(doc,
        ["Control", "Implementación técnica", "Estado"],
        [
            ("TLS 1.2+ obligatorio",       "Azure App Service enforce HTTPS. Solo puerto 443 expuesto.",                   "✅ SI"),
            ("Política de contraseñas",    "Entra External ID B2C: 120 días exp., 9 chars, complejidad, bloqueo 5 intentos.", "✅ SI"),
            ("MFA",                        "Configurable por Entra Conditional Access. SMS o app authenticator.",          "✅ SI"),
            ("CAPTCHA",                    "reCAPTCHA v3 / hCaptcha en login y formularios críticos.",                    "✅ SI"),
            ("Cabeceras de seguridad",     "CSP, HSTS, X-Frame-Options, X-Content-Type-Options. Configurado en Next.js middleware.", "✅ SI"),
            ("WAF",                        "Azure Front Door WAF con reglas OWASP 3.x + DDoS Protection Standard. Nativo Azure, sin dependencia de infraestructura del cliente.", "✅ SI"),
            ("SSO corporativo OIDC",       "Azure Entra ID para 23 usuarios administrativos.",                            "✅ SI"),
            ("Auth local segura",          "Entra External ID B2C para 3,000 colaboradores sin cuenta de red.",           "✅ SI"),
            ("RBAC granular",              "Roles en Next.js API Routes + Row-Level Security en Azure SQL. Scoping por sede.", "✅ SI"),
            ("Cifrado en reposo",          "Azure SQL TDE (AES-256). Azure Blob Storage SSE.",                            "✅ SI"),
            ("Gestión de secretos",        "Azure Key Vault. Cero secrets en código o variables de entorno planas.",      "✅ SI"),
            ("Auditoría completa",         "Log: usuario, acción, timestamp, valores old/new, IP. Exportable.",           "✅ SI"),
            ("Ethical Hacking / Pentest",  "DATTICS ejecuta pentest previo a producción. Resultados entregados a Argos.", "✅ SI"),
            ("Protección de datos",        "Cumplimiento Ley 1581/2012 (Habeas Data Colombia). Finalidad específica definida.", "✅ SI"),
            ("Firma electrónica avanzada", "DocuSign / Certicamara. No repudio, integridad, autenticación, timestamp.",   "✅ SI"),
            ("Retención documental 20 años","Azure Blob Storage con lifecycle management policies. Retención configurable (20+ años). Acceso desde portal admin. Sin dependencia de sistemas externos.", "✅ SI"),
        ]
    )

    h2(doc, "6.2 Certificaciones de plataforma (heredadas de Microsoft Azure)")
    certs = ["ISO 27001", "ISO 27017", "ISO 27018", "SOC 1/2/3 Type II",
             "GDPR", "HIPAA", "CSA STAR", "PCI DSS Level 1"]
    body(doc, "Verificables en el Microsoft Service Trust Portal:")
    p = doc.add_paragraph()
    r = p.add_run("  ·  ".join(certs))
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = C_BLUE
    p.paragraph_format.space_after = Pt(8)

    _page_break(doc)

    # ── 6. COSTOS DE INFRAESTRUCTURA ─────────────────────────────────────────
    h1(doc, "7. Costos de Infraestructura y Comparativa")

    h2(doc, "7.1 Costo mensual de infraestructura Azure")
    table(doc,
        ["Componente", "Especificación", "USD/mes", "COP/mes (aprox.)"],
        [
            ("Azure App Service (Next.js full-stack)", "Standard S2 — Node.js 20 LTS", "$100", "$420,000"),
            ("Azure SQL Database",                 "Standard S2 (50 DTU)",      "$150",  "$630,000"),
            ("Azure AD B2C (Entra External)",      "3,000 autenticaciones/mes", "$15",   "$63,000"),
            ("Azure Functions",                    "Consumption plan (SFTP)",   "$10",   "$42,000"),
            ("Azure Blob Storage",                 "Docs + PDFs + evidencias\n(lifecycle 20 años)", "$15", "$63,000"),
            ("Azure Front Door",                   "Standard (CDN + WAF OWASP)","$35",   "$147,000"),
            ("Azure Application Insights",         "Logs, métricas, alertas",   "$10",   "$42,000"),
            ("Microsoft Graph",                    "Correos vía O365 Argos\n(incluido en licencias M365 del cliente)", "$0", "$0"),
            ("Firma Electrónica Avanzada",         "DocuSign / Certicamara\n(por volumen ~3,000 firmas/año)", "$250", "$1,050,000"),
            ("Azure Key Vault",                    "Gestión de secretos",       "$5",    "$21,000"),
        ]
    )

    body(doc,
        "TOTAL MENSUAL:  ~$590 USD  /  ~$2,478,000 COP\n"
        "Nota: Un único App Service Next.js reemplaza el App Service .NET + Static Web Apps Angular, "
        "reduciendo la complejidad operativa. Microsoft Graph usa el tenant O365 existente de Argos (sin costo adicional). "
        "Power Automate no es parte de esta arquitectura.")

    h2(doc, "7.2 Comparativa de costo total a 3 años (solo licenciamiento)")
    table(doc,
        ["Arquitectura", "Anual (USD)", "3 Años (USD)", "3 Años (COP aprox.)", "Ahorro vs. Power Pages"],
        [
            ("✅ Web Custom + Azure (propuesta)\nSin Power Automate · Sin Power Pages",
             "~$7,188", "~$21,564", "~$91M COP", "—"),
            ("Power Pages autenticado (3,000 users)",
             "$24,000 – $48,000", "$72,000 – $144,000", "~$302M – $605M COP", "Hasta $122,436 USD"),
            ("Power Pages + Power Automate Premium",
             "$48,000 – $96,000", "$144,000 – $288,000", "~$605M – $1,210M COP", "Hasta $266,436 USD"),
        ],
        hdr_bg="1A7A2E"
    )

    body(doc,
        "TRM referencial: 4,200 COP/USD.\n"
        "La propuesta Web Custom elimina por completo las licencias de Power Pages y Power Automate. "
        "La única dependencia de SaaS es la firma electrónica (DocuSign/Certicamara), que también "
        "sería requerida bajo cualquier otra arquitectura.")

    _page_break(doc)

    # ── 7. SLA, DISPONIBILIDAD Y SOPORTE ─────────────────────────────────────
    h1(doc, "8. SLA, Disponibilidad y Modelo de Soporte")

    h2(doc, "8.1 Indicadores de disponibilidad")
    table(doc,
        ["Indicador", "Valor propuesto", "Mecanismo", "SLA Azure"],
        [
            ("Disponibilidad",  "99.9% mensual",
             "Azure App Service slots + Azure SQL geo-replicación",      "99.95% App Service"),
            ("RTO",             "< 4 horas",
             "Azure SQL geo-replicación + App Service deployment slots", "Respaldado financieramente"),
            ("RPO",             "< 1 hora",
             "Backups automáticos Azure SQL cada 5–10 min (retención 35 días)", "99.995% Azure SQL"),
        ]
    )

    h2(doc, "8.2 Modelo de soporte (ANS)")
    table(doc,
        ["Prioridad", "Descripción", "Tiempo respuesta", "Resolución", "Horario"],
        [
            ("P1 – Crítico",  "Sistema caído, compromiso de datos",       "< 1 hora",  "< 4 horas",  "7×24"),
            ("P2 – Alto",     "Función principal afectada, vulnerabilidad","< 2 horas", "< 8 horas",  "L–V 7am–7pm COT"),
            ("P3 – Medio",    "Función secundaria, error no bloqueante",   "< 4 horas", "< 24 horas", "L–V 8am–6pm COT"),
            ("P4 – Bajo",     "Consulta, cambio, mejora menor",           "< 8 horas", "< 72 horas", "L–V 8am–6pm COT"),
        ]
    )

    h2(doc, "8.3 Canales de soporte")
    for c in [
        "ServiceNow de Argos (integración nativa si se requiere).",
        "Correo electrónico de soporte DATTICS.",
        "Canal compartido en Microsoft Teams.",
        "Teléfono para incidentes críticos P1.",
        "Reuniones semanales durante implementación, mensuales durante operación.",
    ]:
        bullet(doc, c)

    _page_break(doc)

    # ── 8. CRONOGRAMA ─────────────────────────────────────────────────────────
    h1(doc, "9. Cronograma de Implementación")

    h2(doc, "Fase 1 — MVP  ·  16 semanas / 4 meses")
    callout(doc,
        "Etapas 1–4 corren en paralelo por pares, reduciendo el calendario sin sacrificar calidad.",
        bg="E8F2FF", border="005287")
    table(doc,
        ["#", "Etapa", "Sem.", "Entregable", "Inicio", "Fin"],
        [
            ("1",  "Levantamiento y diseño funcional detallado",    "2", "HU + wireframes",    "Sem 1",  "Sem 2"),
            ("2",  "Setup infraestructura + CI/CD + ambientes",     "2", "Ambientes OK",        "Sem 1",  "Sem 2"),
            ("3",  "Autenticación SSO + B2C + RBAC",               "3", "Login funcional",     "Sem 3",  "Sem 5"),
            ("4",  "Integración SFTP + SuccessFactors",             "2", "Job SFTP operativo",  "Sem 3",  "Sem 4"),
            ("5",  "Módulo administración (maestros, kits, sedes)", "3", "Admin funcional",     "Sem 5",  "Sem 7"),
            ("6",  "Módulo tallas + ventanas de bloqueo",           "3", "Tallas con ventanas", "Sem 6",  "Sem 8"),
            ("7",  "Registro entrega + firma electrónica + PDF",    "3", "Entrega completa",    "Sem 8",  "Sem 10"),
            ("8",  "Generación SOLPED + cargas masivas",            "2", "SOLPED funcional",    "Sem 10", "Sem 11"),
            ("9",  "Reportes operativos + Power BI",                "2", "Reportes Fase 1",     "Sem 11", "Sem 12"),
            ("10", "Migración datos + pruebas QA + UAT",            "2", "UAT aprobado",        "Sem 12", "Sem 13"),
        ]
    )
    callout(doc,
        "HITO 90 DÍAS  (Sem 13 ≈ 91 días) — Plataforma funcional completa · UAT aprobado · "
        "Módulos 1-10 entregados y validados con el equipo Argos.",
        bg="FFF3E0", border="E36B00")
    table(doc,
        ["#", "Etapa", "Sem.", "Entregable", "Inicio", "Fin"],
        [
            ("11", "Ethical Hacking + remediación de hallazgos",    "2", "Reporte pentest",    "Sem 13", "Sem 14"),
            ("12", "Go-Live + soporte presencial + estabilización", "2", "Producción estable", "Sem 15", "Sem 16"),
        ]
    )
    callout(doc, "TOTAL FASE 1: ~16 semanas (4 meses)", bg="FFF3E0", border="E36B00")

    h2(doc, "Fase 2 — Funcionalidades Avanzadas  ·  6 semanas / ~2 meses")
    table(doc,
        ["#", "Etapa", "Sem.", "Entregable", "Inicio", "Fin"],
        [
            ("13", "Bodegas e inventario en línea",         "3", "Inventario activo",   "Sem 17", "Sem 19"),
            ("14", "Canal PQRS + notificaciones",           "2", "PQRS funcional",       "Sem 19", "Sem 20"),
            ("15", "Alarmas + banners + proveedores",       "2", "Módulos Fase 2",       "Sem 20", "Sem 21"),
            ("16", "Reportes Fase 2 + QA + UAT + Go-Live", "2", "Fase 2 en producción", "Sem 21", "Sem 22"),
        ]
    )
    callout(doc,
        "TOTAL FASE 2: ~6 semanas (2 meses)  |  TOTAL PROYECTO: ~22 semanas (~6 meses)",
        bg="E8F2FF", border="005287")

    _page_break(doc)

    # ── 9. AMBIENTES Y DEVOPS ─────────────────────────────────────────────────
    h1(doc, "10. Ambientes y DevSecOps")

    h2(doc, "10.1 Ambientes")
    table(doc,
        ["Ambiente", "Propósito", "Datos", "Acceso"],
        [
            ("DEV",      "Desarrollo y pruebas unitarias",             "Sintéticos",                       "Equipo DATTICS"),
            ("QA / UAT", "Pruebas funcionales y aceptación del cliente","Subconjunto anonimizado de reales","DATTICS + Equipo Argos"),
            ("PROD",     "Operación productiva",                       "Datos reales",                     "Todos los usuarios"),
        ]
    )

    h2(doc, "10.2 Pipeline CI/CD")
    for step in [
        "Compilación automática Next.js (npm run build) + Azure Functions TypeScript (Azure DevOps).",
        "Pruebas unitarias y análisis de código estático.",
        "Escaneo de dependencias vulnerables (Dependabot).",
        "Escaneo de secretos en código (GitGuardian / GitHub Advanced Security).",
        "Despliegue a slots de staging (blue-green, zero downtime).",
        "Gate de promoción: DEV → QA → PROD con aprobación.",
        "Infrastructure as Code: ARM templates / Terraform.",
    ]:
        bullet(doc, step)

    h2(doc, "10.3 Requisitos de hardware y software para usuarios")
    callout(doc,
        "Ningún requisito adicional. Los usuarios solo necesitan un dispositivo con navegador web moderno "
        "(Chrome, Edge, Firefox, Safari). Sin instalación de software cliente, agentes, plugins ni Java. "
        "La solución es 100% cloud — no requiere On-Premises Gateway.",
        bg="E8F2FF", border="005287"
    )

    _page_break(doc)

    # ── 10. RIESGOS Y MITIGACIONES ────────────────────────────────────────────
    h1(doc, "11. Riesgos y Mitigaciones")

    table(doc,
        ["Riesgo", "Impacto", "Prob.", "Mitigación"],
        [
            ("Disponibilidad del SFTP de SuccessFactors",
             "Alto", "Media",
             "Job con reintentos exponenciales. Log de errores visible para Admin. Ejecución retroactiva manual."),
            ("Conectividad con SAP para API directa (futuro)",
             "Medio", "Baja",
             "Fase 1 usa archivo plano (ya validado). Arquitectura preparada para API/BAPI sin cambiar portal."),
            ("Volumen de firma electrónica (picos de entrega)",
             "Medio", "Media",
             "Pool de licencias DocuSign/Certicamara con volumen acordado. Escalamiento elástico."),
            ("Cambios en formato de archivos SSFF",
             "Alto", "Baja",
             "Parser configurable en Azure Functions. Alertas automáticas si el formato cambia."),
            ("Adopción de colaboradores (3,000 usuarios)",
             "Alto", "Media",
             "UX simplificada tipo móvil. Plan de capacitación por sede. Soporte N1 desde mesa Argos."),
            ("Volumen de almacenamiento a largo plazo (20 años de PDFs)",
             "Bajo", "Baja",
             "Azure Blob Storage con lifecycle policies. Tier Hot → Cool → Archive automático por antigüedad. Costo decrece con el tiempo."),
        ]
    )

    # ── 11. MATRIZ DE TRAZABILIDAD ────────────────────────────────────────────
    h1(doc, "12. Matriz de Trazabilidad — Requerimientos RFP")

    table(doc,
        ["Requerimiento RFP (Anexo 03 / 04)", "Cobertura en la propuesta", "Fase"],
        [
            ("SSO corporativo (Azure AD / OIDC)",
             "Entra ID OIDC para 23 usuarios administrativos",                         "MVP"),
            ("Auth colaboradores sin cuenta de red",
             "Entra External ID B2C: cédula + contraseña + MFA",                       "MVP"),
            ("4 perfiles de rol con permisos específicos",
             "RBAC en Next.js API Routes + Row-Level Security Azure SQL",              "MVP"),
            ("Actualización de tallas self-service",
             "Ventanas de bloqueo parametrizables. Colaborador actualiza desde portal", "MVP"),
            ("Motor de kits por cargo/género/sede/ciclo",
             "Servicio TypeScript interno en Next.js. Recalculo automático al cambiar posición/sede.", "MVP"),
            ("Registro de entrega con firma electrónica avanzada (Ley 527/1999)",
             "DocuSign / Certicamara. PDF generado server-side por Next.js. Envío vía Microsoft Graph. Almacenado en Azure Blob Storage.", "MVP"),
            ("Generación archivo SOLPED para SAP",
             "Módulo Pedido Masivo. Archivo plano con formato exacto SAP",             "MVP"),
            ("Integración SuccessFactors",
             "Azure Functions SFTP diario: altas, bajas, cambios cargo/sede",          "MVP"),
            ("Carga masiva de dotación y tallas por sede",
             "Módulo de cargas masivas con validación Excel/CSV",                      "MVP"),
            ("Reportes operativos + export Excel/PDF",
             "7 reportes en Fase 1, dashboards Power BI en Fase 2",                   "MVP / F2"),
            ("Administración de bodegas e inventario",
             "Bodegas por sede, movimientos, consolidado nacional",                    "Fase 2"),
            ("Canal de PQRS y garantías",
             "Portal PQRS bidireccional, notificación automática proveedor",           "Fase 2"),
            ("Devoluciones y garantías",
             "Registradas como tipo de PQRS con flujo especializado",                 "Fase 2"),
            ("TLS 1.2+, WAF, CSP/HSTS",
             "Azure Front Door WAF (OWASP 3.x) + DDoS Standard. Next.js middleware con headers CSP, HSTS, X-Frame-Options.", "MVP"),
            ("Retención documental 20 años",
             "Azure Blob Storage con lifecycle management policies (Hot → Cool → Archive). Sin dependencia de sistemas externos.", "MVP"),
            ("Pentesting obligatorio pre-producción",
             "DATTICS ejecuta Ethical Hacking y entrega informe a Argos",             "MVP"),
            ("SLA 99.9%, RTO < 4h, RPO < 1h",
             "Azure SQL geo-replicación + App Service slots",                         "MVP"),
        ]
    )

    _page_break(doc)

    # ── 12. EQUIPO TÉCNICO ────────────────────────────────────────────────────
    h1(doc, "13. Equipo Técnico del Proyecto")

    callout(doc,
        "El proyecto cuenta con un equipo multidisciplinario de especialistas en desarrollo "
        "web, arquitectura cloud Azure, integración de sistemas empresariales y seguridad. "
        "Cada rol tiene dedicación definida según la fase del proyecto.",
        bg="E8F2FF", border="005287")

    h2(doc, "13.1 Estructura del equipo")
    table(doc,
        ["Rol", "Perfil / Experiencia requerida", "Responsabilidades principales", "Dedicación", "Fase"],
        [
            ("Líder de Proyecto (PM)",
             "PMP o equivalente. Experiencia en proyectos de transformación digital con clientes "
             "corporativos en Colombia. Conocimiento de metodologías ágiles (Scrum/Kanban). "
             "Gestión de stakeholders y ANS.",
             "Planificación y seguimiento del cronograma. Gestión de riesgos y cambios. "
             "Coordinación con stakeholders Argos. Reporte de avance semanal. "
             "Cumplimiento de ANS contractuales.",
             "100%", "Todas"),
            ("Arquitecto de Solución (SA)",
             "Especialista en arquitectura cloud Azure (certificación AZ-204 / AZ-305 deseable). "
             "Diseño de sistemas distribuidos, seguridad por diseño, integración de APIs REST y "
             "OIDC. Experiencia con Next.js / Node.js en entornos enterprise.",
             "Diseño de arquitectura técnica Next.js full-stack y decisiones de patrones. "
             "Guía técnica al equipo. Revisión de seguridad y cumplimiento. "
             "Validación de integraciones y contratos API. Alineación con estándares Microsoft.",
             "50%", "Todas"),
            ("Desarrollador Full-Stack Senior (Next.js — UI)",
             "4+ años en Next.js 13+ / React 18. TypeScript avanzado. App Router, Server Components, "
             "Server Actions. Diseño de componentes accesibles, responsive (mobile-first). "
             "Fluent UI React. Conocimiento de WCAG 2.1.",
             "Arquitectura del frontend Next.js y design system. Módulos críticos (autenticación "
             "MSAL, tallas, entregas, admin). Revisión de código UI. Garantía de performance y "
             "accesibilidad WCAG 2.1.",
             "100%", "F1 + F2"),
            ("Desarrollador Full-Stack Semi-Senior (Next.js — UI)",
             "2+ años en Next.js / React. TypeScript. Desarrollo de páginas y componentes, "
             "formularios con React Hook Form, integración con API Routes, Zustand o Context API.",
             "Desarrollo de módulos UI (tallas, reportes, notificaciones, PQRS). "
             "Formularios reactivos y validaciones client-side. Consumo de API Routes. "
             "Pruebas unitarias con Jest + Testing Library.",
             "100%", "F1 + F2"),
            ("Desarrollador Full-Stack Senior (Next.js — API / Datos)",
             "4+ años en Next.js / Node.js. TypeScript avanzado. API Routes, Server Actions, "
             "Prisma ORM, Azure SQL. Patrones repositorio, autorización por claims. "
             "Experiencia con integraciones SFTP y generación de archivos planos para ERP.",
             "Diseño de API Routes y Server Actions. Motor de kits y generador de SOLPED "
             "(TypeScript). Seguridad de endpoints (RBAC, scoping). Modelado Prisma + Azure SQL.",
             "100%", "F1 + F2"),
            ("Desarrollador Full-Stack Semi-Senior (Next.js — API / Datos)",
             "2+ años en Node.js / TypeScript. API Routes, lógica de negocio, validaciones "
             "con Zod, pruebas unitarias (Vitest / Jest). Soporte a integraciones.",
             "Desarrollo de endpoints secundarios y lógica de negocio. Validaciones con Zod. "
             "Pruebas unitarias (Vitest). Soporte al especialista de integraciones.",
             "100%", "F1 + F2"),
            ("Especialista en Integraciones",
             "Experiencia con Azure Functions (TypeScript/Node.js, timer trigger), "
             "SFTP client en Node.js, Microsoft Graph API (send mail, O365). "
             "Conocimiento de generación de archivos SOLPED para SAP y conectores REST "
             "de firma electrónica (DocuSign / Certicamara).",
             "Job SFTP para sincronización con SuccessFactors. Generación y envío de archivos "
             "SOLPED a SAP. Integración Microsoft Graph para notificaciones O365. "
             "Conectores de firma electrónica (DocuSign / Certicamara).",
             "60%", "F1 + F2"),
            ("Ingeniero DevOps / Azure",
             "Certificación AZ-400 o equivalente. Azure DevOps, CI/CD pipelines, "
             "ARM Templates / Terraform, Azure App Service (Node.js), Key Vault, "
             "Application Insights. Blue-green deployment.",
             "Provisión de infraestructura Azure (IaC: ARM / Terraform). Pipelines CI/CD "
             "con build Next.js y gate de promoción. Gestión de secretos en Key Vault. "
             "Monitoreo con Application Insights. Blue-green deployment y rollback.",
             "50%", "Todas"),
            ("Diseñador UX/UI",
             "Diseño de interfaces para aplicaciones empresariales en entornos industriales "
             "(planta, bodega). Wireframing, prototipado, design system. Experiencia con "
             "Fluent UI React o shadcn/ui. Enfoque mobile-first.",
             "Wireframes y prototipos de alta fidelidad. Design system en Fluent UI React. "
             "Experiencia mobile-first para usuarios en planta y bodega. Validación de "
             "accesibilidad y usabilidad con el cliente.",
             "70%", "F1"),
            ("Analista Funcional / QA",
             "Levantamiento de historias de usuario, casos de prueba, validación de reglas de "
             "negocio. Pruebas funcionales, de regresión y UAT con el cliente. Documentación "
             "de manuales de usuario y SOPs.",
             "Levantamiento y refinamiento de historias de usuario. Casos de prueba y escenarios "
             "UAT. Validación de reglas de negocio con Argos. Documentación de manuales de "
             "usuario y SOPs operativos.",
             "100%", "Todas"),
            ("Especialista en Ciberseguridad / Ethical Hacking",
             "Certificación OSCP, CEH o equivalente. Pruebas de penetración (OWASP Top 10), "
             "revisión de cabeceras HTTP, análisis de tokens JWT, escaneo de dependencias. "
             "Entrega de informe formal a Argos previo a Go-Live.",
             "Pruebas de penetración (OWASP Top 10) sobre la plataforma. Revisión de "
             "configuración Azure (Key Vault, Front Door, CORS). Análisis de tokens JWT y "
             "cabeceras HTTP. Informe formal de seguridad entregado a Argos previo a Go-Live.",
             "30% (Pentest)", "Pre-Go-Live"),
        ]
    )

    h2(doc, "13.2 Matriz RACI del equipo")
    table(doc,
        ["Actividad", "PM", "SA", "FS UI", "FS API", "DevOps", "QA / Analista"],
        [
            ("Diseño de arquitectura y HU",               "A", "R", "C", "C", "C", "C"),
            ("Setup infraestructura Azure y CI/CD",       "A", "C", "I", "I", "R", "I"),
            ("Desarrollo módulos UI (React / Next.js)",   "A", "C", "R", "C", "I", "C"),
            ("Desarrollo API Routes / Server Actions",    "A", "C", "C", "R", "I", "C"),
            ("Integraciones (SSFF, SAP, O365)",           "A", "C", "I", "C", "C", "C"),
            ("Pruebas unitarias y de integración",        "A", "I", "R", "R", "C", "C"),
            ("UAT con el cliente (Argos)",                "A", "C", "C", "C", "I", "R"),
            ("Pentest y Ethical Hacking",                 "A", "C", "I", "I", "C", "C"),
            ("Capacitación y documentación",              "A", "I", "C", "C", "I", "R"),
            ("Go-Live y soporte post-lanzamiento",        "A", "C", "C", "C", "R", "C"),
        ],
        hdr_bg="005287"
    )

    body(doc, "R = Responsable de ejecutar   A = Aprobador final   C = Consultado   I = Informado")

    h2(doc, "13.3 Dedicación por fase")
    table(doc,
        ["Rol", "Fase 0\n(Discovery)", "Fase 1\n(MVP — 16 sem.)", "Fase 2\n(Avanzada — 6 sem.)", "Soporte\n(post Go-Live)"],
        [
            ("Líder de Proyecto",            "100%", "100%", "100%", "50%"),
            ("Arquitecto de Solución",       "100%", "50%",  "30%",  "20%"),
            ("FS Senior — UI (Next.js)",     "50%",  "100%", "100%", "30%"),
            ("FS Semi-Senior — UI (Next.js)","30%",  "100%", "100%", "20%"),
            ("FS Senior — API (Next.js)",    "50%",  "100%", "100%", "30%"),
            ("FS Semi-Senior — API (Next.js)","30%", "100%", "80%",  "20%"),
            ("Especialista Integraciones",   "80%",  "60%",  "60%",  "20%"),
            ("DevOps / Azure",               "100%", "50%",  "30%",  "20%"),
            ("Diseñador UX/UI",              "100%", "70%",  "30%",  "0%"),
            ("Analista Funcional / QA",      "100%", "100%", "100%", "50%"),
            ("Especialista Ciberseguridad",  "0%",   "30%",  "0%",   "0%"),
        ]
    )

    _page_break(doc)

    # ── 13. CONCLUSIÓN ────────────────────────────────────────────────────────
    h1(doc, "14. Conclusión Ejecutiva")

    callout(doc,
        "La arquitectura Web Custom sobre Azure es la opción más eficiente, segura y económica "
        "para digitalizar la gestión de dotación de Cementos Argos S.A.S. Ofrece el mismo nivel "
        "de seguridad enterprise de Microsoft Azure, con control total sobre la lógica de negocio "
        "compleja y a un costo de hasta 7 veces menor que Power Pages + Power Automate Premium.",
        bg="E8F2FF", border="005287"
    )

    body(doc,
        "La propuesta de DATTICS S.A.S. entrega:\n\n"
        "  •  Una plataforma self-service para los 3,000+ colaboradores: desde su celular o PC, "
        "sin instalar nada, pueden actualizar tallas, consultar dotación, firmar entregas y "
        "registrar PQRS.\n\n"
        "  •  Herramientas potentes para los 4 perfiles administrativos: cargas masivas, "
        "gestión de inventario/bodegas, control de ventanas de bloqueo, generación automática "
        "del archivo SAP.\n\n"
        "  •  Código fuente como activo de Argos, desplegado en su propio tenant Azure, sin "
        "dependencia de licencias perpetuas de Power Platform.\n\n"
        "  •  Cumplimiento del 100% de los 85 requerimientos técnicos del Anexo 04 y los "
        "Lineamientos de Ciberseguridad del Grupo Argos.")

    doc.add_paragraph()
    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_firma.add_run(
        "DATTICS S.A.S. — Equipo de Arquitectura de Solución\n"
        "Marzo 2026  ·  Versión 3.0  ·  Confidencial")
    r.font.size = Pt(9)
    r.font.color.rgb = C_GRAY
    r.italic = True

    return doc


# ─── Entrada ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    out = "/Users/eduardosacahui/Github-Repos/ARG_Plataforma_Integral_Dotacion/ARGOS_Propuesta_WebCustom_RFP.docx"
    build().save(out)
    print(f"✅  Documento generado: {out}")
